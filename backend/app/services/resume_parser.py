import io
import logging
from typing import Dict, Any, Optional
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document

from app.core.ai_service import ai_service

logger = logging.getLogger(__name__)

class ResumeParserService:
    """
    Service for handling resume uploads and parsing
    Extracts text from PDF/DOCX and uses AI to structure the data
    """
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self.ai_service = ai_service
    
    async def parse_resume_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        Main method to parse resume file and extract structured data
        
        Args:
            file: Uploaded resume file (PDF or DOCX)
            
        Returns:
            Structured profile data extracted by AI
        """
        try:
            # Validate file
            self._validate_file(file)
            
            # Extract text from file
            resume_text = await self._extract_text_from_file(file)
            
            if not resume_text.strip():
                raise HTTPException(
                    status_code=422, 
                    detail="Could not extract text from resume file"
                )
            
            # Use AI to parse the resume text
            parsed_data = await self.ai_service.parse_resume(resume_text)
            
            # Generate AI insights
            insights = await self.ai_service.generate_profile_insights(parsed_data)
            
            # Combine parsed data with insights
            result = {
                **parsed_data,
                "ai_insights": insights,
                "raw_text_length": len(resume_text),
                "parsing_success": True
            }
            
            logger.info(f"Successfully parsed resume with {len(resume_text)} characters")
            return result
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error parsing resume file: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Error processing resume: {str(e)}"
            )
    
    def _validate_file(self, file: UploadFile) -> None:
        """
        Validate uploaded file format and size
        """
        if not file.filename:
            raise HTTPException(status_code=422, detail="No file provided")
        
        # Check file extension
        file_extension = self._get_file_extension(file.filename)
        if file_extension not in self.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=422,
                detail=f"Unsupported file type. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}"
            )
        
        # Check file size (if available)
        if hasattr(file, 'size') and file.size and file.size > self.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=422,
                detail=f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB"
            )
    
    async def _extract_text_from_file(self, file: UploadFile) -> str:
        """
        Extract text content from PDF or DOCX file
        """
        file_extension = self._get_file_extension(file.filename)
        
        # Read file content
        content = await file.read()
        
        if file_extension == '.pdf':
            return self._extract_text_from_pdf(content)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_text_from_docx(content)
        else:
            raise HTTPException(status_code=422, detail="Unsupported file format")
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF file using PyPDF2
        """
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise HTTPException(
                status_code=422,
                detail="Could not read PDF file. Please ensure it's a valid PDF."
            )
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX file using python-docx
        """
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise HTTPException(
                status_code=422,
                detail="Could not read DOCX file. Please ensure it's a valid Word document."
            )
    
    def _get_file_extension(self, filename: str) -> str:
        """
        Get file extension from filename
        """
        return '.' + filename.split('.')[-1].lower() if '.' in filename else ''
    
    async def update_profile_from_text(self, resume_text: str) -> Dict[str, Any]:
        """
        Parse resume from plain text (for testing or direct text input)
        """
        try:
            parsed_data = await self.ai_service.parse_resume(resume_text)
            insights = await self.ai_service.generate_profile_insights(parsed_data)
            
            return {
                **parsed_data,
                "ai_insights": insights,
                "raw_text_length": len(resume_text),
                "parsing_success": True
            }
            
        except Exception as e:
            logger.error(f"Error parsing resume text: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Error processing resume text: {str(e)}"
            )

# Global resume parser instance
resume_parser = ResumeParserService() 